from docx import Document
from docx.shared import Pt
document = Document('path/to/file.docx')

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)
"""
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']
"""
for paragraph in document.paragraphs:
    paragraph.style = document.styles['Normal']
    for run in paragraph.runs:
        run.font.size = Pt(10)    
document.save('refactored.docx')
